from __future__ import annotations

import html
import re
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import httpx
from docx import Document
from pypdf import PdfReader

from humanizer.core.errors import ValidationError

CODE_SUFFIXES = {
    ".c",
    ".cc",
    ".cpp",
    ".cs",
    ".go",
    ".java",
    ".js",
    ".jsx",
    ".kt",
    ".php",
    ".py",
    ".rb",
    ".rs",
    ".scala",
    ".sh",
    ".sql",
    ".swift",
    ".ts",
    ".tsx",
}


def resolve_text_input(
    text: str | None,
    input_path: str | None,
    input_url: str | None = None,
) -> str:
    if text and text.strip():
        return text.strip()
    if input_path:
        return load_text_from_path(input_path)
    if input_url:
        return load_text_from_url(input_url)
    raise ValidationError("one of text, input_path, or input_url must be provided")


def load_text_from_path(input_path: str) -> str:
    path = Path(input_path).expanduser()
    if not path.exists():
        raise ValidationError(f"input file does not exist: {input_path}")
    if not path.is_file():
        raise ValidationError(f"input path is not a file: {input_path}")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"} or suffix in CODE_SUFFIXES:
        return _read_plain_text(path)
    if suffix == ".pdf":
        return _read_pdf_text(path)
    if suffix == ".docx":
        return _read_docx_text(path)
    raise ValidationError(f"unsupported input type: {suffix or 'no_extension'}")


def _read_plain_text(path: Path) -> str:
    content = path.read_text(encoding="utf-8").strip()
    if not content:
        raise ValidationError(f"input file is empty: {path}")
    return content


def _read_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    content = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    if not content:
        raise ValidationError(f"no extractable text found in pdf: {path}")
    return content


def _read_docx_text(path: Path) -> str:
    document = Document(str(path))
    content = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()).strip()
    if not content:
        raise ValidationError(f"no extractable text found in docx: {path}")
    return content


def load_text_from_url(input_url: str) -> str:
    response = httpx.get(input_url, follow_redirects=True, timeout=10.0)
    response.raise_for_status()

    parsed = urlparse(input_url)
    suffix = Path(parsed.path).suffix.lower()
    content_type = response.headers.get("content-type", "").split(";")[0].strip().lower()

    if suffix in {".txt", ".md", ".markdown"} or suffix in CODE_SUFFIXES:
        return _normalize_remote_text(response.text, input_url)
    if suffix == ".pdf" or content_type == "application/pdf":
        return _read_pdf_blob(response.content, input_url)
    if suffix == ".docx" or content_type == DOCX_CONTENT_TYPE:
        return _read_docx_blob(response.content, input_url)
    if content_type in {"text/plain", "text/markdown"}:
        return _normalize_remote_text(response.text, input_url)
    if content_type in {"text/html", "application/xhtml+xml"} or not content_type:
        return _extract_html_text(response.text, input_url)
    raise ValidationError(f"unsupported url content type: {content_type or suffix or 'unknown'}")


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _normalize_remote_text(content: str, source: str) -> str:
    normalized = content.strip()
    if not normalized:
        raise ValidationError(f"no text content found at url: {source}")
    return normalized


def _read_pdf_blob(blob: bytes, source: str) -> str:
    reader = PdfReader(BytesIO(blob))
    content = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    if not content:
        raise ValidationError(f"no extractable text found in pdf: {source}")
    return content


def _read_docx_blob(blob: bytes, source: str) -> str:
    document = Document(BytesIO(blob))
    content = "\n".join(
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ).strip()
    if not content:
        raise ValidationError(f"no extractable text found in docx: {source}")
    return content


def _extract_html_text(markup: str, source: str) -> str:
    text = re.sub(r"<script.*?>.*?</script>", " ", markup, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<style.*?>.*?</style>", " ", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    text = " ".join(text.split()).strip()
    if not text:
        raise ValidationError(f"no text content found at url: {source}")
    return text


def detect_content_type(
    text: str,
    input_path: str | None,
    requested_content_type: str,
) -> str:
    if requested_content_type != "auto":
        return requested_content_type
    suffix = Path(input_path or "").suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".pdf", ".docx"}:
        return "text"
    if suffix in CODE_SUFFIXES:
        return "code"
    if "```" in text:
        return "text"
    code_markers = ["def ", "class ", "function ", "import ", "#include", "const ", "let ", "var "]
    if sum(1 for marker in code_markers if marker in text) >= 2:
        return "code"
    return "text"
